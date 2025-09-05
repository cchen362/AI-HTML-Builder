import os
import tempfile
from typing import Optional
from pathlib import Path
from datetime import datetime
import aiofiles
import structlog
from docx import Document
from fastapi import UploadFile, HTTPException
from ..core.config import settings

logger = structlog.get_logger()

class FileProcessor:
    MAX_SIZE = settings.max_upload_size  # 50MB
    ALLOWED_EXTENSIONS = ['.txt', '.md', '.docx']
    ALLOWED_MIME_TYPES = [
        'text/plain',
        'text/markdown',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    ]
    
    @staticmethod
    def validate_file(file: UploadFile) -> dict:
        """Enhanced file validation with detailed error reporting"""
        validation_result = {
            'valid': True,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        try:
            # Basic file info
            file_size = getattr(file, 'size', 0) if hasattr(file, 'size') else 0
            file_name = file.filename or 'unknown'
            content_type = file.content_type or 'unknown'
            
            validation_result['file_info'] = {
                'name': file_name,
                'size': file_size,
                'content_type': content_type,
                'size_mb': round(file_size / (1024 * 1024), 2) if file_size > 0 else 0
            }
            
            # Check if filename exists
            if not file.filename:
                validation_result['errors'].append("No filename provided")
                validation_result['valid'] = False
            
            # Check file size
            if file_size > FileProcessor.MAX_SIZE:
                max_size_mb = round(FileProcessor.MAX_SIZE / (1024 * 1024), 1)
                current_size_mb = round(file_size / (1024 * 1024), 1)
                validation_result['errors'].append(
                    f"File too large: {current_size_mb}MB (max: {max_size_mb}MB)"
                )
                validation_result['valid'] = False
            elif file_size == 0:
                validation_result['warnings'].append("File appears to be empty")
            
            # Check file extension
            if file.filename:
                file_ext = Path(file.filename).suffix.lower()
                validation_result['file_info']['extension'] = file_ext
                
                if not file_ext:
                    validation_result['errors'].append("File has no extension")
                    validation_result['valid'] = False
                elif file_ext not in FileProcessor.ALLOWED_EXTENSIONS:
                    validation_result['errors'].append(
                        f"Unsupported file type '{file_ext}'. Allowed types: {', '.join(FileProcessor.ALLOWED_EXTENSIONS)}"
                    )
                    validation_result['valid'] = False
            
            # Check MIME type
            if file.content_type:
                if file.content_type not in FileProcessor.ALLOWED_MIME_TYPES:
                    validation_result['warnings'].append(
                        f"MIME type '{file.content_type}' not in standard allowed list"
                    )
                    
                    # Check for common mismatches
                    if file.content_type == 'application/octet-stream':
                        validation_result['warnings'].append(
                            "Generic binary MIME type detected - file type detection may be unreliable"
                        )
            else:
                validation_result['warnings'].append("No MIME type provided")
            
            # Enhanced logging
            logger.info(
                "File validation completed",
                filename=file_name,
                valid=validation_result['valid'],
                size_mb=validation_result['file_info']['size_mb'],
                extension=validation_result['file_info'].get('extension', 'none'),
                errors_count=len(validation_result['errors']),
                warnings_count=len(validation_result['warnings'])
            )
            
            if not validation_result['valid']:
                error_details = "; ".join(validation_result['errors'])
                raise HTTPException(
                    status_code=400, 
                    detail=f"File validation failed: {error_details}"
                )
            
            return validation_result
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error("File validation failed unexpectedly", filename=file.filename, error=str(e))
            raise HTTPException(status_code=500, detail=f"File validation error: {str(e)}")
    
    @staticmethod
    async def process_file(file: UploadFile) -> str:
        """
        Enhanced file processing with comprehensive error handling
        
        Supports:
        - .txt files (plain text)
        - .md files (markdown)  
        - .docx files (Word documents)
        """
        validation_result = None
        try:
            # Enhanced validation with detailed reporting
            validation_result = FileProcessor.validate_file(file)
            
            if not file.filename:
                raise HTTPException(status_code=400, detail="No filename provided")
            
            file_ext = validation_result['file_info'].get('extension', '')
            file_size = validation_result['file_info'].get('size', 0)
            
            logger.info(
                "Starting file processing",
                filename=file.filename,
                extension=file_ext,
                size_mb=validation_result['file_info'].get('size_mb', 0),
                warnings=validation_result.get('warnings', [])
            )
            
            # Read file content with error handling
            try:
                content = await file.read()
                if not content:
                    raise HTTPException(status_code=400, detail="File is empty or could not be read")
                    
                logger.debug(f"Read {len(content)} bytes from file: {file.filename}")
            except Exception as e:
                logger.error("Failed to read file content", filename=file.filename, error=str(e))
                raise HTTPException(status_code=400, detail="Could not read file content")
            
            # Reset file pointer for potential reuse
            try:
                await file.seek(0)
            except Exception as e:
                logger.warning("Could not reset file pointer", filename=file.filename, error=str(e))
            
            # Process based on file type
            extracted_text = ""
            try:
                if file_ext in ['.txt', '.md']:
                    extracted_text = await FileProcessor._process_text_file(content, file.filename)
                elif file_ext == '.docx':
                    extracted_text = await FileProcessor._process_docx_file(content, file.filename)
                else:
                    raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")
            except HTTPException:
                raise
            except Exception as e:
                logger.error("File content processing failed", filename=file.filename, file_type=file_ext, error=str(e))
                raise HTTPException(status_code=500, detail=f"Failed to process {file_ext} file: {str(e)}")
            
            # Validate extracted content
            if not extracted_text or not extracted_text.strip():
                raise HTTPException(
                    status_code=400, 
                    detail=f"No readable text content found in {file.filename}. The file may be empty, corrupted, or contain only non-text elements."
                )
            
            logger.info(
                "File processing completed successfully",
                filename=file.filename,
                extracted_length=len(extracted_text),
                word_count=len(extracted_text.split()) if extracted_text else 0
            )
            
            return extracted_text
                
        except HTTPException:
            raise
        except Exception as e:
            error_context = {
                'filename': getattr(file, 'filename', 'unknown'),
                'validation_info': validation_result['file_info'] if validation_result else None,
                'error': str(e)
            }
            logger.error("Unexpected file processing failure", **error_context)
            raise HTTPException(
                status_code=500, 
                detail=f"Unexpected error processing file: {str(e)}"
            )
    
    @staticmethod
    async def _process_text_file(content: bytes, filename: str) -> str:
        """Enhanced text and markdown file processing with encoding detection"""
        try:
            if not content:
                raise ValueError("Empty content provided")
            
            # Try different encodings in order of preference
            encodings = [
                'utf-8',      # Most common
                'utf-16',     # Common for Windows files
                'utf-16-le',  # Little-endian UTF-16
                'utf-16-be',  # Big-endian UTF-16
                'latin-1',    # ISO-8859-1
                'cp1252',     # Windows-1252
                'ascii'       # Basic ASCII
            ]
            
            last_error = None
            for encoding in encodings:
                try:
                    text_content = content.decode(encoding)
                    
                    # Basic content validation
                    if len(text_content.strip()) == 0:
                        logger.warning(f"File decoded successfully with {encoding} but contains only whitespace", filename=filename)
                    
                    # Check for replacement characters (indicates encoding issues)
                    if '\ufffd' in text_content and encoding != 'utf-8':
                        logger.warning(f"Replacement characters found with {encoding} encoding", filename=filename)
                        continue
                    
                    logger.info("Successfully decoded text file", 
                               filename=filename, 
                               encoding=encoding,
                               content_length=len(text_content),
                               lines=len(text_content.splitlines()))
                    
                    return text_content.strip()
                    
                except UnicodeDecodeError as e:
                    last_error = e
                    logger.debug(f"Failed to decode with {encoding}", filename=filename, error=str(e))
                    continue
            
            # If all specific encodings fail, use utf-8 with error handling
            logger.warning("All preferred encodings failed, using UTF-8 with error replacement", 
                          filename=filename, last_error=str(last_error))
            
            text_content = content.decode('utf-8', errors='replace')
            
            # Count replacement characters to assess quality
            replacement_count = text_content.count('\ufffd')
            if replacement_count > 0:
                logger.warning(f"Used fallback UTF-8 encoding with {replacement_count} replacement characters", 
                              filename=filename)
            
            if len(text_content.strip()) == 0:
                raise ValueError("File contains no readable text content after decoding")
            
            return text_content.strip()
            
        except ValueError:
            raise
        except Exception as e:
            logger.error("Text file processing failed", filename=filename, error=str(e))
            raise ValueError(f"Could not process text file: {str(e)}")
    
    @staticmethod
    async def _process_docx_file(content: bytes, filename: str) -> str:
        """Enhanced Word document processing with comprehensive error handling"""
        tmp_file_path = None
        try:
            if not content:
                raise ValueError("Empty DOCX content provided")
            
            # Validate that content looks like a DOCX file (ZIP format)
            if not content.startswith(b'PK'):
                raise ValueError("File does not appear to be a valid DOCX file (missing ZIP header)")
            
            # Create temporary file with error handling
            try:
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    tmp_file.write(content)
                    tmp_file_path = tmp_file.name
                    
                logger.debug(f"Created temporary DOCX file: {tmp_file_path}", filename=filename)
            except Exception as e:
                logger.error("Failed to create temporary file for DOCX processing", filename=filename, error=str(e))
                raise ValueError(f"Could not create temporary file: {str(e)}")
            
            try:
                # Load document with python-docx
                try:
                    doc = Document(tmp_file_path)
                except Exception as e:
                    logger.error("Failed to load DOCX document", filename=filename, error=str(e))
                    
                    # Provide more specific error messages
                    error_msg = str(e).lower()
                    if 'bad zip file' in error_msg or 'not a zip file' in error_msg:
                        raise ValueError("File appears to be corrupted or is not a valid DOCX file")
                    elif 'permission' in error_msg:
                        raise ValueError("Permission denied while processing DOCX file")
                    else:
                        raise ValueError(f"Could not open DOCX file: {str(e)}")
                
                # Extract text from all paragraphs
                text_parts = []
                paragraph_count = 0
                
                try:
                    for paragraph in doc.paragraphs:
                        paragraph_count += 1
                        text = paragraph.text.strip()
                        if text:  # Only add non-empty paragraphs
                            text_parts.append(text)
                except Exception as e:
                    logger.error("Error extracting paragraphs from DOCX", filename=filename, error=str(e))
                    raise ValueError(f"Could not extract paragraphs: {str(e)}")
                
                # Extract text from tables
                table_count = 0
                try:
                    for table in doc.tables:
                        table_count += 1
                        for row_idx, row in enumerate(table.rows):
                            row_text = []
                            for cell_idx, cell in enumerate(row.cells):
                                try:
                                    cell_text = cell.text.strip()
                                    if cell_text:
                                        row_text.append(cell_text)
                                except Exception as e:
                                    logger.warning(f"Error extracting cell text (table {table_count}, row {row_idx}, cell {cell_idx})", 
                                                  filename=filename, error=str(e))
                                    continue
                            
                            if row_text:
                                text_parts.append(" | ".join(row_text))
                except Exception as e:
                    logger.warning("Error extracting tables from DOCX (continuing with paragraphs only)", 
                                  filename=filename, error=str(e))
                
                # Combine all extracted text
                extracted_text = "\n\n".join(text_parts)
                
                # Validate extraction results
                if not extracted_text.strip():
                    raise ValueError("No readable text content found in DOCX file. The document may be empty or contain only images/objects.")
                
                # Log successful processing with statistics
                word_count = len(extracted_text.split()) if extracted_text else 0
                logger.info(
                    "Successfully processed DOCX",
                    filename=filename,
                    paragraphs=paragraph_count,
                    tables=table_count,
                    text_parts=len(text_parts),
                    final_length=len(extracted_text),
                    word_count=word_count
                )
                
                return extracted_text
                
            finally:
                # Clean up temporary file with error handling
                if tmp_file_path:
                    try:
                        os.unlink(tmp_file_path)
                        logger.debug(f"Cleaned up temporary file: {tmp_file_path}")
                    except Exception as e:
                        logger.warning("Could not delete temporary DOCX file", 
                                      temp_path=tmp_file_path, error=str(e))
                    
        except ValueError:
            raise
        except Exception as e:
            logger.error("Unexpected DOCX processing failure", filename=filename, error=str(e))
            raise ValueError(f"Unexpected error processing DOCX file: {str(e)}")
    
    @staticmethod
    async def save_temp_file(content: str, session_id: str, filename: str) -> str:
        """Save processed content to temporary file"""
        try:
            # Create temp directory structure
            temp_dir = Path(tempfile.gettempdir()) / "ai-html-builder" / "uploads" / session_id
            temp_dir.mkdir(parents=True, exist_ok=True)
            
            # Create unique filename
            timestamp = int(datetime.utcnow().timestamp())
            safe_filename = f"{timestamp}_{filename}.txt"
            temp_file_path = temp_dir / safe_filename
            
            # Save content
            async with aiofiles.open(temp_file_path, 'w', encoding='utf-8') as f:
                await f.write(content)
            
            logger.info("Temp file saved", path=str(temp_file_path))
            return str(temp_file_path)
            
        except Exception as e:
            logger.error("Failed to save temp file", session_id=session_id, filename=filename, error=str(e))
            raise
    
    @staticmethod
    async def cleanup_session_files(session_id: str):
        """Clean up temporary files for a session"""
        try:
            session_dir = Path(tempfile.gettempdir()) / "ai-html-builder" / "uploads" / session_id
            if session_dir.exists():
                import shutil
                shutil.rmtree(session_dir)
                logger.info("Session files cleaned up", session_id=session_id)
        except Exception as e:
            logger.error("Failed to cleanup session files", session_id=session_id, error=str(e))
    
    @staticmethod
    def get_file_info(file: UploadFile) -> dict:
        """Get file information"""
        return {
            "name": file.filename or "unknown",
            "size": getattr(file, 'size', 0),
            "type": file.content_type or "unknown"
        }