"""File processing utilities for uploaded documents."""

from __future__ import annotations

import asyncio
import csv
import io
from typing import Any

import structlog

logger = structlog.get_logger()

ALLOWED_EXTENSIONS: frozenset[str] = frozenset(
    {".txt", ".md", ".docx", ".pdf", ".csv", ".xlsx"}
)


class FileProcessingError(Exception):
    """Raised when file validation or processing fails."""


def validate_file(filename: str, file_size: int, max_size_mb: int = 50) -> None:
    """Validate file extension and size. Raises FileProcessingError on failure."""
    max_bytes = max_size_mb * 1024 * 1024
    if file_size > max_bytes:
        raise FileProcessingError(
            f"File size ({file_size:,} bytes) exceeds maximum of {max_size_mb}MB"
        )
    ext = _get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise FileProcessingError(
            f"File type '{ext}' not allowed. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )


def _get_extension(filename: str) -> str:
    """Extract lowercase extension including the dot."""
    if "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


async def process_file(filename: str, content: bytes) -> dict[str, Any]:
    """Process an uploaded file and return extracted content.

    Runs synchronous file parsing in a thread to avoid blocking the event loop.
    """
    return await asyncio.to_thread(_process_file_sync, filename, content)


def _process_file_sync(filename: str, content: bytes) -> dict[str, Any]:
    """Synchronous file processing implementation.

    Returns dict with keys:
        filename, file_type, content_type ("text"|"data"),
        content, data_preview, row_count, columns
    """
    ext = _get_extension(filename)

    if ext in {".txt", ".md"}:
        text = _process_text(content)
        return _text_result(filename, ext, text)

    if ext == ".docx":
        text = _process_docx(content)
        return _text_result(filename, ext, text)

    if ext == ".pdf":
        text = _process_pdf(content)
        return _text_result(filename, ext, text)

    if ext == ".csv":
        rows, columns = _process_csv(content)
        preview = (
            f"CSV Data: {len(rows)} rows, {len(columns)} columns\n"
            f"Columns: {', '.join(columns)}"
        )
        return _data_result(filename, ext, rows, columns, preview)

    if ext == ".xlsx":
        rows, columns, sheet_name = _process_xlsx(content)
        preview = (
            f"Excel Data: {len(rows)} rows, {len(columns)} columns\n"
            f"Sheet: {sheet_name}\nColumns: {', '.join(columns)}"
        )
        return _data_result(filename, ext, rows, columns, preview)

    raise FileProcessingError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _text_result(filename: str, ext: str, text: str) -> dict[str, Any]:
    return {
        "filename": filename,
        "file_type": ext,
        "content_type": "text",
        "content": text,
        "data_preview": None,
        "row_count": None,
        "columns": None,
    }


def _data_result(
    filename: str,
    ext: str,
    rows: list[dict[str, str]],
    columns: list[str],
    preview: str,
) -> dict[str, Any]:
    return {
        "filename": filename,
        "file_type": ext,
        "content_type": "data",
        "content": _rows_to_text(rows, columns, max_rows=20),
        "data_preview": preview,
        "row_count": len(rows),
        "columns": columns,
    }


def _process_text(content: bytes) -> str:
    """Process .txt or .md files."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1")


def _process_docx(content: bytes) -> str:
    """Extract text from .docx file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                tables_text.append(row_text)

    all_text = "\n\n".join(paragraphs)
    if tables_text:
        all_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)
    return all_text


def _process_pdf(content: bytes) -> str:
    """Extract text from .pdf file."""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page_num, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        if page_text and page_text.strip():
            parts.append(f"--- Page {page_num} ---\n{page_text}")
    return "\n\n".join(parts)


def _process_csv(content: bytes) -> tuple[list[dict[str, str]], list[str]]:
    """Parse .csv file into rows and column names."""
    text = content.decode("utf-8-sig")  # Handle BOM
    reader = csv.DictReader(io.StringIO(text))
    rows = list(reader)
    columns = list(rows[0].keys()) if rows else list(reader.fieldnames or [])
    return rows, columns


def _process_xlsx(
    content: bytes,
) -> tuple[list[dict[str, str]], list[str], str]:
    """Parse .xlsx file into rows, columns, and sheet name."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    sheet = wb.active
    if sheet is None:
        raise FileProcessingError("Excel workbook has no active sheet")
    sheet_name: str = sheet.title or "Sheet1"

    headers: list[str] = []
    for cell in sheet[1]:
        headers.append(str(cell.value) if cell.value is not None else "")

    rows: list[dict[str, str]] = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        row_dict: dict[str, str] = {}
        for header, value in zip(headers, row):
            row_dict[header] = str(value) if value is not None else ""
        rows.append(row_dict)

    wb.close()
    return rows, headers, sheet_name


def _rows_to_text(
    rows: list[dict[str, str]], columns: list[str], max_rows: int = 20
) -> str:
    """Convert rows of data to a readable text summary."""
    lines: list[str] = []
    header_line = " | ".join(columns)
    lines.append(header_line)
    lines.append("-" * len(header_line))
    for row in rows[:max_rows]:
        lines.append(" | ".join(row.get(c, "") for c in columns))
    if len(rows) > max_rows:
        lines.append(f"... ({len(rows) - max_rows} more rows)")
    return "\n".join(lines)


def generate_upload_prompt(file_result: dict[str, Any]) -> str:
    """Generate an appropriate chat prompt from upload results."""
    filename = file_result["filename"]
    content_type = file_result["content_type"]
    content = file_result["content"]

    # Truncate to 2000 chars for the prompt
    preview = content[:2000]
    if len(content) > 2000:
        preview += "\n... [full content provided]"

    if content_type == "data":
        data_preview = file_result.get("data_preview", "")
        return (
            f"Create an interactive dashboard from this data:\n\n"
            f"File: {filename}\n{data_preview}\n\n{preview}"
        )

    ext = _get_extension(filename)
    if ext in {".txt", ".md"}:
        return (
            f"Create a styled HTML document from this content:\n\n"
            f"File: {filename}\n\n{preview}"
        )
    if ext in {".docx", ".pdf"}:
        return (
            f"Create a professional document from this content:\n\n"
            f"File: {filename}\n\n{preview}"
        )

    return f"Create a document from this content:\n\nFile: {filename}\n\n{preview}"
